"""Extract plain text from uploaded documents (.pdf, .docx, .txt)."""

import io


class UnsupportedDocumentError(ValueError):
    pass


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext == "docx":
        return _extract_docx(file_bytes)
    if ext in ("txt", "md", "rtf"):
        return file_bytes.decode("utf-8", errors="replace")
    raise UnsupportedDocumentError(f"Формат .{ext} не поддерживается. Пришли PDF, DOCX или TXT.")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
